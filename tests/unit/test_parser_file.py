"""Tests for parser file-type routing and Excel/Word parsing via temporary files."""

from __future__ import annotations

import json
import os
import tempfile
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from electronics_rfq_agent.parser import RFQParser

SAMPLE_JSON_STR = json.dumps(
    [
        {
            "line_number": 1,
            "part_number": "RES-0402-10K",
            "quantity": 100,
            "required_date": None,
            "manufacturer": "Yageo",
            "customer_notes": None,
        },
        {
            "line_number": 2,
            "part_number": "CAP-100NF",
            "quantity": 50,
            "required_date": None,
            "manufacturer": None,
            "customer_notes": "blue tape",
        },
    ]
)


@pytest.fixture
def parser() -> RFQParser:
    return RFQParser(model="claude-sonnet-4-6")


class TestParseRouting:
    @pytest.mark.asyncio
    async def test_parse_routes_to_text_for_nonexistent_path(
        self, parser: RFQParser
    ) -> None:
        """String that isn't a real file path should go to _parse_text."""
        mock_response = MagicMock()
        mock_response.content = [MagicMock(text=SAMPLE_JSON_STR)]
        mock_client = MagicMock()
        mock_client.messages.create = AsyncMock(return_value=mock_response)

        with patch("anthropic.AsyncAnthropic", return_value=mock_client):
            items = await parser.parse("not a real file path xyz")
        assert len(items) == 2

    @pytest.mark.asyncio
    async def test_parse_routes_txt_file_to_text(self, parser: RFQParser) -> None:
        """Existing .txt file should be read and passed to _parse_text."""
        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False) as f:
            f.write("RES-0402-10K, 100\nCAP-100NF, 50")
            tmp_path = f.name

        try:
            mock_response = MagicMock()
            mock_response.content = [MagicMock(text=SAMPLE_JSON_STR)]
            mock_client = MagicMock()
            mock_client.messages.create = AsyncMock(return_value=mock_response)

            with patch("anthropic.AsyncAnthropic", return_value=mock_client):
                items = await parser.parse(tmp_path)
            assert len(items) == 2
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_routes_xlsx_to_excel(self, parser: RFQParser) -> None:
        """Existing .xlsx file should route to _parse_excel."""
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None
            ws.append(["Part Number", "Quantity", "Manufacturer"])
            ws.append(["RES-0402-10K", 100, "Yageo"])
            ws.append(["CAP-100NF", 50, "Murata"])
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 2
            assert items[0].part_number == "RES-0402-10K"
            assert items[0].quantity == 100
            assert items[1].quantity == 50
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xlsx_manufacturer_captured(self, parser: RFQParser) -> None:
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None
            ws.append(["Part Number", "Quantity", "Manufacturer", "Notes"])
            ws.append(["RES-001", 10, "Yageo", "Priority"])
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert items[0].manufacturer == "Yageo"
            assert items[0].customer_notes == "Priority"
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xlsx_skips_empty_rows(self, parser: RFQParser) -> None:
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None
            ws.append(["Part Number", "Quantity"])
            ws.append(["RES-001", 10])
            ws.append([None, None])  # empty row — should be skipped
            ws.append(["CAP-001", 20])
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 2
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xlsx_quantity_defaults_to_1(self, parser: RFQParser) -> None:
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None
            ws.append(["Part Number", "Quantity"])
            ws.append(["RES-001", None])  # None quantity
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 1
            assert items[0].quantity == 1
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xlsx_empty_workbook_returns_empty(
        self, parser: RFQParser
    ) -> None:
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            # No rows at all — _find_header_row returns None
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert items == []
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xlsx_no_recognizable_header_returns_empty(
        self, parser: RFQParser
    ) -> None:
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None
            # Rows with no recognizable header keywords
            ws.append(["col_x", "col_y", "col_z"])
            ws.append(["data1", "data2", "data3"])
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert items == []
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_routes_docx_to_word(self, parser: RFQParser) -> None:
        """Existing .docx file should route to _parse_word."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
        try:
            doc = Document()
            table = doc.add_table(rows=3, cols=3)
            table.cell(0, 0).text = "Part Number"
            table.cell(0, 1).text = "Quantity"
            table.cell(0, 2).text = "Manufacturer"
            table.cell(1, 0).text = "RES-0402-10K"
            table.cell(1, 1).text = "100"
            table.cell(1, 2).text = "Yageo"
            table.cell(2, 0).text = "CAP-100NF"
            table.cell(2, 1).text = "50"
            table.cell(2, 2).text = "Murata"
            doc.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 2
            assert items[0].part_number == "RES-0402-10K"
            assert items[0].quantity == 100
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_skips_table_without_part_header(
        self, parser: RFQParser
    ) -> None:
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
        try:
            doc = Document()
            # Table without recognizable headers
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "col_a"
            table.cell(0, 1).text = "col_b"
            table.cell(1, 0).text = "data1"
            table.cell(1, 1).text = "data2"
            doc.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert items == []
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_empty_rows_skipped(self, parser: RFQParser) -> None:
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
        try:
            doc = Document()
            table = doc.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "part number"
            table.cell(0, 1).text = "qty"
            table.cell(1, 0).text = "RES-001"
            table.cell(1, 1).text = "10"
            table.cell(2, 0).text = ""  # empty row
            table.cell(2, 1).text = ""
            doc.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 1
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_pdf_calls_anthropic(self, parser: RFQParser) -> None:
        """PDF path should call Anthropic with base64 content."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", mode="wb", delete=False) as f:
            f.write(b"%PDF-1.4 fake pdf content")
            tmp_path = f.name

        try:
            mock_response = MagicMock()
            mock_response.content = [MagicMock(text=SAMPLE_JSON_STR)]
            mock_client = MagicMock()
            mock_client.messages.create = AsyncMock(return_value=mock_response)

            with patch("anthropic.AsyncAnthropic", return_value=mock_client):
                items = await parser.parse(tmp_path)

            assert len(items) == 2
            mock_client.messages.create.assert_called_once()
            call_kwargs = mock_client.messages.create.call_args.kwargs
            # Should pass the document as base64
            messages = call_kwargs.get("messages", [])
            assert len(messages) == 1
            content = messages[0]["content"]
            assert any(
                item.get("type") == "document"
                for item in content
                if isinstance(item, dict)
            )
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xlsx_multi_sheet_finds_bom_sheet(
        self, parser: RFQParser
    ) -> None:
        """If the active sheet has no BOM header, parser searches other sheets."""
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            # Active (first) sheet = cover page with no recognizable BOM header
            summary_ws = wb.active
            assert summary_ws is not None
            summary_ws.title = "Summary"
            summary_ws.append(["Quote Date", "2026-06-19"])
            summary_ws.append(["Customer", "ACME Corp"])

            # Second sheet = BOM with line items
            bom_ws = wb.create_sheet("BOM")
            bom_ws.append(["Part Number", "Quantity", "Manufacturer"])
            bom_ws.append(["RES-0402-10K", 100, "Yageo"])
            bom_ws.append(["CAP-100NF", 50, "Murata"])
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 2
            assert items[0].part_number == "RES-0402-10K"
            assert items[1].part_number == "CAP-100NF"
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_pdf_raises_rfq_parse_error_on_empty_content(
        self, parser: RFQParser
    ) -> None:
        """RFQParseError is raised when Anthropic returns empty content list."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", mode="wb", delete=False) as f:
            f.write(b"%PDF-1.4 fake content")
            tmp_path = f.name
        try:
            mock_response = MagicMock()
            mock_response.content = []
            mock_client = MagicMock()
            mock_client.messages.create = AsyncMock(return_value=mock_response)

            from electronics_rfq_agent.models import RFQParseError

            with patch("anthropic.AsyncAnthropic", return_value=mock_client):
                with pytest.raises(RFQParseError, match="no parseable text"):
                    await parser.parse(tmp_path)
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_text_raises_rfq_parse_error_on_empty_content(
        self, parser: RFQParser
    ) -> None:
        """RFQParseError raised when Anthropic returns no text for plain-text input."""
        mock_response = MagicMock()
        mock_response.content = []
        mock_client = MagicMock()
        mock_client.messages.create = AsyncMock(return_value=mock_response)

        from electronics_rfq_agent.models import RFQParseError

        with patch("anthropic.AsyncAnthropic", return_value=mock_client):
            with pytest.raises(RFQParseError, match="no parseable text"):
                await parser.parse("RES-001, qty 10")

    @pytest.mark.asyncio
    async def test_parse_xlsx_fractional_quantity(self, parser: RFQParser) -> None:
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = f.name
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None
            ws.append(["Part Number", "Quantity"])
            ws.append(["RES-001", 10.5])
            wb.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 1
            assert items[0].quantity == 10
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_docx_fractional_quantity(self, parser: RFQParser) -> None:
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
        try:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Part Number"
            table.cell(0, 1).text = "Quantity"
            table.cell(1, 0).text = "RES-001"
            table.cell(1, 1).text = "10.5"
            doc.save(tmp_path)

            items = await parser.parse(tmp_path)
            assert len(items) == 1
            assert items[0].quantity == 10
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_binary_txt_raises_rfq_parse_error(
        self, parser: RFQParser
    ) -> None:
        from electronics_rfq_agent.models import RFQParseError

        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"\x00\xff\xfe binary")
            tmp_path = f.name
        try:
            with pytest.raises(RFQParseError):
                await parser.parse(tmp_path)
        finally:
            os.unlink(tmp_path)

    @pytest.mark.asyncio
    async def test_parse_xls_routes_to_excel(self, parser: RFQParser) -> None:
        """A .xls suffix should also route to _parse_excel."""
        import openpyxl

        with tempfile.NamedTemporaryFile(suffix=".xls", delete=False) as f:
            tmp_path = f.name

        try:
            with patch.object(parser, "_parse_excel", wraps=parser._parse_excel) as spy:
                wb = openpyxl.Workbook()
                ws = wb.active
                assert ws is not None
                ws.append(["Part Number", "Quantity"])
                ws.append(["DIODE-001", 25])
                wb.save(tmp_path)

                try:
                    items = await parser.parse(tmp_path)
                except Exception:
                    pytest.skip(
                        "openpyxl does not support .xls extension on this platform"
                    )

                spy.assert_called_once()
                assert isinstance(items, list)
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
